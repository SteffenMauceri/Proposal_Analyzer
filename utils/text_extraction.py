from pathlib import Path
from typing import Optional
import re
import sys

# For PDF text extraction
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

# For Word document text extraction
try:
    import docx
except ImportError:
    docx = None

# For legacy .doc files (optional, requires python-docx2txt)
try:
    import docx2txt
except ImportError:
    docx2txt = None

def extract_text_from_pdf(pdf_path: Path) -> Optional[str]:
    """
    Extracts text content from a PDF file.
    
    Args:
        pdf_path (Path): The path to the PDF file.
        
    Returns:
        Optional[str]: The extracted text, or None if extraction fails or PyPDF2 is not available.
    """
    if PyPDF2 is None:
        print("Warning: PyPDF2 library is not installed. PDF text extraction will not be available.", file=sys.stderr)
        return None
        
    if not pdf_path.is_file():
        return None
    
    try:
        raw_text_parts = []
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    raw_text_parts.append(page_text)
        
        if not raw_text_parts:
            return None

        full_raw_text = "\n".join(raw_text_parts)
        
        # 1. Normalize Unicode and handle mojibake (basic)
        # This is a complex topic. For now, we assume text is mostly okay or rely on LLM for minor issues.
        # A more robust solution might involve libraries like ftfy.
        # For now, let's focus on structural cleanup.

        # 2. Re-join hyphenated words at line breaks
        # This regex looks for a word character, a hyphen, a newline, and another word character.
        processed_text = re.sub(r'(\w)-(\r\n|\r|\n)(\w)', r'\1\3', full_raw_text)
        
        # 3. Normalize all forms of newlines to a single \n, then handle multiple newlines
        processed_text = re.sub(r'(\r\n|\r|\n)+', '\n', processed_text)
        
        # 4. Normalize other whitespace (multiple spaces/tabs to single space)
        # but preserve newlines for paragraph structure if they are meaningful
        lines = processed_text.split('\n')
        cleaned_lines = [re.sub(r'[ \t\xA0]+', ' ', line).strip() for line in lines]
        # Rejoin lines. If paragraphs are important, consider double newline, but LLM prompt will handle it.
        processed_text = "\n".join(cleaned_lines)
        
        # 5. Remove excessive blank lines (more than 2 consecutive newlines)
        processed_text = re.sub(r'\n{3,}', '\n\n', processed_text)
        
        return processed_text.strip() # Final strip for any leading/trailing whitespace

    except Exception as e:
        # Optionally log the error e
        print(f"Error extracting text from PDF {pdf_path}: {e}", file=sys.stderr)
        return None

def extract_text_from_docx(docx_path: Path) -> Optional[str]:
    """
    Extracts text content from a .docx file.
    
    Args:
        docx_path (Path): The path to the .docx file.
        
    Returns:
        Optional[str]: The extracted text, or None if extraction fails or python-docx is not available.
    """
    if docx is None:
        print("Warning: python-docx library is not installed. DOCX text extraction will not be available.", file=sys.stderr)
        return None
        
    if not docx_path.is_file():
        return None
    
    try:
        doc = docx.Document(docx_path)
        full_text = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        if not full_text:
            return None
            
        raw_text = "\n".join(full_text)
        
        # Apply similar cleaning as PDF extraction
        # 1. Normalize newlines
        processed_text = re.sub(r'(\r\n|\r|\n)+', '\n', raw_text)
        
        # 2. Normalize whitespace
        lines = processed_text.split('\n')
        cleaned_lines = [re.sub(r'[ \t\xA0]+', ' ', line).strip() for line in lines]
        processed_text = "\n".join(cleaned_lines)
        
        # 3. Remove excessive blank lines
        processed_text = re.sub(r'\n{3,}', '\n\n', processed_text)
        
        return processed_text.strip()
        
    except Exception as e:
        print(f"Error extracting text from DOCX {docx_path}: {e}", file=sys.stderr)
        return None

def extract_text_from_doc(doc_path: Path) -> Optional[str]:
    """
    Extracts text content from a legacy .doc file.
    
    Args:
        doc_path (Path): The path to the .doc file.
        
    Returns:
        Optional[str]: The extracted text, or None if extraction fails or docx2txt is not available.
    """
    if docx2txt is None:
        print("Warning: docx2txt library is not installed. Legacy DOC text extraction will not be available.", file=sys.stderr)
        print("Install with: pip install docx2txt", file=sys.stderr)
        return None
        
    if not doc_path.is_file():
        return None
    
    try:
        # docx2txt.process() can handle both .doc and .docx files
        raw_text = docx2txt.process(str(doc_path))
        
        if not raw_text or not raw_text.strip():
            return None
            
        # Apply similar cleaning as other extraction methods
        # 1. Normalize newlines
        processed_text = re.sub(r'(\r\n|\r|\n)+', '\n', raw_text)
        
        # 2. Normalize whitespace
        lines = processed_text.split('\n')
        cleaned_lines = [re.sub(r'[ \t\xA0]+', ' ', line).strip() for line in lines]
        processed_text = "\n".join(cleaned_lines)
        
        # 3. Remove excessive blank lines
        processed_text = re.sub(r'\n{3,}', '\n\n', processed_text)
        
        return processed_text.strip()
        
    except Exception as e:
        print(f"Error extracting text from DOC {doc_path}: {e}", file=sys.stderr)
        return None

def extract_text_from_document(document_path: Path) -> Optional[str]:
    """
    Extracts text content from a document file (PDF, DOCX, or DOC).
    
    Args:
        document_path (Path): The path to the document file.
        
    Returns:
        Optional[str]: The extracted text, or None if extraction fails.
    """
    if not document_path.is_file():
        return None
    
    file_extension = document_path.suffix.lower()
    
    if file_extension == '.pdf':
        return extract_text_from_pdf(document_path)
    elif file_extension == '.docx':
        return extract_text_from_docx(document_path)
    elif file_extension == '.doc':
        return extract_text_from_doc(document_path)
    else:
        print(f"Warning: Unsupported file format '{file_extension}' for text extraction from {document_path}", file=sys.stderr)
        return None 